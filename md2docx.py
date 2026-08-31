import re, sys, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn

URL_RE = re.compile(r'(https?://[^\s)]+|(?:read\.)?levantermarkets\.com)')

def add_hyperlink(p, url, text):
    part = p.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement('w:hyperlink'); h.set(qn('r:id'), r_id)
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '1155CC'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    r.append(rPr); t = OxmlElement('w:t'); t.text = text; r.append(t); h.append(r)
    p._p.append(h)

def add_fmt(p, text):
    i = 0
    for m in re.finditer(r'\*\*(.+?)\*\*|\*(.+?)\*', text):
        if m.start() > i: p.add_run(text[i:m.start()])
        if m.group(1) is not None:
            run = p.add_run(m.group(1)); run.bold = True
        else:
            run = p.add_run(m.group(2)); run.italic = True
        i = m.end()
    if i < len(text): p.add_run(text[i:])

def add_runs(p, text):
    pos = 0
    for m in URL_RE.finditer(text):
        add_fmt(p, text[pos:m.start()])
        url = m.group(0); href = url if url.startswith('http') else 'https://' + url
        add_hyperlink(p, href, url); pos = m.end()
    add_fmt(p, text[pos:])

SEP_RE = re.compile(r'^\|[\s:|-]+\|$')


def _cells(s):
    """Split a markdown table row into its cells."""
    return [c.strip() for c in s.strip().strip('|').split('|')]


def flush_table(doc, rows):
    """Render buffered markdown table rows as a real Word table. Without this a
    table line lands in the document as a paragraph full of pipe characters."""
    if not rows:
        return
    body = [r for r in rows if not SEP_RE.match(r)]
    if not body:
        return
    grid = [_cells(r) for r in body]
    width = max(len(r) for r in grid)
    t = doc.add_table(rows=0, cols=width)
    t.style = 'Table Grid'
    for i, row in enumerate(grid):
        cells = t.add_row().cells
        for j in range(width):
            cell = cells[j]
            cell.text = ''
            add_runs(cell.paragraphs[0], row[j] if j < len(row) else '')
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def convert(md, out):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal'].font.size = Pt(11)
    tbuf = []
    for raw in open(md, encoding='utf-8').read().split('\n'):
        s = raw.strip()
        if s.startswith('|') and s.endswith('|') and len(s) > 1:
            tbuf.append(s); continue
        if tbuf:
            flush_table(doc, tbuf); tbuf = []
        if not s or s == '>': continue
        quote = False
        if s.startswith('> '): quote = True; s = s[2:].strip()
        if s == '---':
            doc.add_paragraph(); continue
        if s.startswith('# '): doc.add_heading(s[2:], level=0); continue
        if s.startswith('## '): doc.add_heading(s[3:], level=1); continue
        if s.startswith('### '): doc.add_heading(s[4:], level=2); continue
        if s.startswith('- ') or s.startswith('* '):
            p = doc.add_paragraph(style='List Bullet'); add_runs(p, s[2:])
        else:
            p = doc.add_paragraph(); add_runs(p, s)
        if quote: p.paragraph_format.left_indent = Inches(0.3)
    flush_table(doc, tbuf)
    doc.save(out)
    return out


def convert_dir(src_dir):
    """Convert every .md in src_dir to src_dir/docx/*.docx."""
    import glob
    ddir = os.path.join(src_dir, "docx")
    os.makedirs(ddir, exist_ok=True)
    n = 0
    for md in sorted(glob.glob(os.path.join(src_dir, "*.md"))):
        base = os.path.splitext(os.path.basename(md))[0]
        convert(md, os.path.join(ddir, base + ".docx"))
        n += 1
    return n


if __name__ == "__main__":
    import sys
    d = sys.argv[1] if len(sys.argv) > 1 else "reports/substack"
    print(f"converted {convert_dir(d)} files to {d}/docx/")
